from docx import Document
doc = Document(r"e:\Hackathon\LuminFlow\docs\demo\LuminFlow_Hackathon_Submission.docx")
print(f"Tables: {len(doc.tables)}")
img_count = sum(1 for r in doc.part.rels.values() if "image" in r.reltype)
print(f"Images: {img_count}")
print(f"Sections: {len(doc.sections)}")
s = doc.sections[0]
print(f"Page: {s.page_width} x {s.page_height}")
print(f"Margins: T={s.top_margin} B={s.bottom_margin} L={s.left_margin} R={s.right_margin}")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t and i >= 73 and i <= 125:
        print(f"[{i:03d}] [{p.style.name}] {t[:250]}")

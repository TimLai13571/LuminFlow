from docx import Document
doc = Document(r'e:\Hackathon\LuminFlow\docs\demo\LuminFlow_Hackathon_Submission.docx')
print(f'Total paragraphs: {len(doc.paragraphs)}')
print()
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f'[{i:03d}] [{p.style.name}] {t[:150]}')

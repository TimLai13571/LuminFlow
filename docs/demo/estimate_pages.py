from docx import Document
from docx.shared import Pt, Emu, Inches
import re

doc = Document(r"e:\Hackathon\LuminFlow\docs\demo\LuminFlow_Hackathon_Submission.docx")

# Page dimensions
s = doc.sections[0]
page_height_emu = s.page_height  # EMU
top_margin = s.top_margin or 914400
bottom_margin = s.bottom_margin or 914400
usable_height_emu = page_height_emu - top_margin - bottom_margin

# Convert to points (1 pt = 12700 EMU)
usable_height_pt = usable_height_emu / 12700

# Font size map from styles
default_font_pt = 10  # We set this in styles
heading_fonts = {"Heading 1": 15, "Heading 2": 12, "Heading 3": 11}
title_font = 24

total_height_pt = 0

for p in doc.paragraphs:
    t = p.text.strip()
    if not t:
        total_height_pt += 6  # empty paragraph
        continue
    
    style = p.style.name if p.style else "Normal"
    font_pt = heading_fonts.get(style, default_font_pt)
    if style == "Title":
        font_pt = title_font
    
    # Line spacing: 1.0 for body, headings have more
    line_spacing = 1.0
    if "Heading" in style:
        line_spacing = 1.15
    
    line_height = font_pt * line_spacing
    
    # Count lines (rough estimate: chars per line ~90 for A4 with 1" margins)
    chars_per_line = 85 if style == "List Paragraph" else 90
    text_len = len(t)
    num_lines = max(1, (text_len + chars_per_line - 1) // chars_per_line)
    
    # Add spacing before/after for headings
    extra = 0
    if style == "Heading 1":
        extra = 12  # ~8pt before + 4pt after after reduction
    elif style == "Heading 2":
        extra = 8
    elif style == "Heading 3":
        extra = 6
    elif style == "Title":
        extra = 14
    
    total_height_pt += num_lines * line_height + extra

# Add table height
for table in doc.tables:
    total_height_pt += len(table.rows) * 14  # rough per-row height

# Add image height
total_height_pt += 300  # rough image height in pt

pages = total_height_pt / usable_height_pt
print(f"Total content height: {total_height_pt:.0f} pt")
print(f"Usable page height: {usable_height_pt:.0f} pt")
print(f"Estimated pages: {pages:.1f}")
print(f"Rounded up: {int(pages + 0.99)} pages")

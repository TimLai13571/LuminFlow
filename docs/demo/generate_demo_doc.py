from pathlib import Path
from shutil import copyfile
from PIL import Image

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"e:\Hackathon\LuminFlow")
DEMO_DIR = ROOT / "docs" / "demo"
SCREEN_DIR = DEMO_DIR / "screenshots"
SOURCE_DIR = Path(r"C:\Program Files\Qoder")
OUTPUT = DEMO_DIR / "LuminFlow_Platform_Demo_For_Consultants.docx"

SCREENSHOTS = {
    "dashboard": "luminflow-dashboard.png",
    "objective": "luminflow-objective.png",
    "sampling": "luminflow-sampling.png",
    "impact_config": "luminflow-impact-config.png",
    "impact_loading": "luminflow-impact-loading.png",
    "impact_result": "luminflow-impact-result.png",
    "team": "luminflow-team.png",
    "role": "luminflow-role-switcher.png",
    "ai": "luminflow-ai-drawer.png",
    "language": "luminflow-language-zh.png",
}

PAGES = [
    {
        "title": "Dashboard",
        "screenshot": "dashboard",
        "goal": "Serves as the audit project cockpit, centrally presenting project progress, key metrics, risk hotspots, control status, pending items, and activity timeline.",
        "modules": [
            "KPI Cards: Filter by role to display overall audit progress, control coverage, PBC completion rate, deficiency count, and risk scores.",
            "Risk Heatmap: Renders risk-control hotspots based on the RCM matrix and RAWTC scores, helping quickly identify high-risk control points.",
            "Status Doughnut / Pending Items: Shows control test status distribution, overdue PBC items, and pending tasks.",
            "Activity Timeline: Records key deliverables, approvals, client responses, and review activities.",
        ],
        "flow": [
            "Enter the Dashboard to review overall project health and key risk indicators.",
            "Switch between CFO, Audit Manager, and Partner views to observe how the same project appears with different permission scopes.",
            "Navigate from high-risk KPIs, overdue PBCs, or the risk heatmap into Objective, Sampling, or Team modules for deeper analysis.",
        ],
        "value": "The Dashboard enables management and project teams to communicate progress and risks on a shared factual basis, reducing information asymmetry in audit status reporting.",
        "feedback": "Consultants should assess whether KPI definitions cover the three core needs: project management, quality review, and client communication.",
    },
    {
        "title": "Audit Objectives",
        "screenshot": "objective",
        "goal": "Visualizes relationships between audit objectives, risk areas, control points, test procedures, and findings using a tree diagram and linked cards.",
        "modules": [
            "AuditTreeD3: Displays the audit objective hierarchy, control point statuses, deficiency nodes, and risk areas.",
            "Node Detail: Click a node to view control description, execution status, risk score, findings, and remediation suggestions.",
            "RadarChart / DualProgress: Shows COSO coverage, audit progress, and client-side response progress.",
            "FindingsTable / AIInsightPanel: Summarizes deficiency severity, deviation rates, impact assessments, and AI priority recommendations.",
        ],
        "flow": [
            "Locate a control point or finding using the search box, e.g., CTRL-002 Income Proof Review.",
            "Click a node to open details and inspect test procedures, COSO mapping, and deficiency evidence.",
            "Combine the radar chart, dual progress bars, and findings table to determine the next review or remediation priority.",
        ],
        "value": "The Objectives page transforms the traditional audit matrix into an explorable graph, making it easier for consultants to inspect test coverage completeness and control-deficiency propagation relationships.",
        "feedback": "Consultants should evaluate whether node granularity, COSO mapping, and deficiency priorities align with actual audit methodology.",
    },
    {
        "title": "Sampling Preview",
        "screenshot": "sampling",
        "goal": "Generates explainable smart sampling plans based on population size, risk scores, confidence levels, prior-year deficiencies, and system changes.",
        "modules": [
            "SampleSizeCard: Displays recommended sample range, population size, confidence level, and period.",
            "ParamAdjuster: Allows adjustment of risk threshold, confidence level, and high-risk control weighting factor.",
            "SunburstChart / TimeWindowCard: Shows document category distribution and time-window distribution.",
            "AILogicExplainer / SampleTable: Explains the sampling rationale and displays sample details.",
        ],
        "flow": [
            "Consultant reviews the system-recommended sample range and document distribution.",
            "Adjust confidence level or risk threshold, then recalculate sample size.",
            "Review the AI explanation, historical trends, and sample details to determine whether the sampling plan is auditable and defensible under challenge.",
        ],
        "value": "Smart Sampling transforms sampling rationale from black-box experience into explainable, adjustable, and auditable risk-oriented plans.",
        "feedback": "Consultants should assess whether sampling parameters, sample distribution, and explanation text are sufficient to support workpaper documentation.",
    },
    {
        "title": "Impact Simulation",
        "screenshot": "impact_result",
        "goal": "Simulates the cascading impact of control deficiencies, organizational changes, regulatory changes, or IT system changes on audit scope, sample size, timeline, and risk indicators.",
        "modules": [
            "EventConfig: Configure event type, severity, impact dimensions, and response strategy.",
            "ForceGraphD3: Displays direct, indirect, and potential impact nodes as a network graph.",
            "ComparisonTable: Compares residual risk, control effectiveness, compliance indicators, operational efficiency, and cost impact before and after simulation.",
            "AIRecommendations: Generates mitigation suggestions, priorities, and risk reduction estimates based on simulation results.",
        ],
        "flow": [
            "Select a change event and set severity, impact dimensions, and response strategy.",
            "Click Run Simulation and wait for the network node propagation animation and computation to complete.",
            "Review impact chains, indicator changes, and AI recommendations, then decide whether to adjust audit scope or add compensating controls.",
        ],
        "value": "Impact Simulation transforms risk events from static descriptions into dynamic decision-support evidence, enabling consultants to proactively assess the impact of changes on audit plans.",
        "feedback": "Consultants should focus on verifying whether impact dimensions, propagation logic, and AI recommendations align with professional audit judgment.",
    },
    {
        "title": "Team Management",
        "screenshot": "team",
        "goal": "Supports transparent collaboration among audit teams, clients, and partners through visibility management, AI approval workflows, client interaction statistics, and evidence checking.",
        "modules": [
            "VisibilityPanel: Controls content boundaries for client-visible, team-internal, and partner-review views.",
            "ApprovalQueue: Routes AI-generated content through Audit Manager approval and Partner QC review, preventing unconfirmed information from being released externally.",
            "InteractionChart / HotQuestionsBar: Tracks client questions, hot topics, and response efficiency.",
            "EvidenceChecker: Checks critical evidence for missing, overdue, and review status.",
        ],
        "flow": [
            "Audit Manager reviews team collaboration and visibility status.",
            "Approve AI-generated summaries or client responses, escalating to Partner review when necessary.",
            "Combine client interaction hotspots with evidence check results to drive PBC supplementation and deficiency closure.",
        ],
        "value": "The Team page increases transparency while preserving audit quality control and permission boundaries, reducing risks in client communication and AI output.",
        "feedback": "Consultants should evaluate whether approval granularity, evidence check rules, and client visibility boundaries meet project quality control requirements.",
    },
]


def ensure_dirs():
    DEMO_DIR.mkdir(parents=True, exist_ok=True)
    SCREEN_DIR.mkdir(parents=True, exist_ok=True)


def crop_for_doc(src: Path, dst: Path):
    with Image.open(src) as image:
        image = image.convert("RGB")
        width, height = image.size
        crop_height = min(height, 1100)
        cropped = image.crop((0, 0, width, crop_height))
        if cropped.width > 1800:
            ratio = 1800 / cropped.width
            cropped = cropped.resize((1800, int(cropped.height * ratio)), Image.Resampling.LANCZOS)
        cropped.save(dst, "PNG", optimize=True)


def prepare_screenshots():
    copied = {}
    for key, name in SCREENSHOTS.items():
        source = SOURCE_DIR / name
        if not source.exists():
            raise FileNotFoundError(f"Missing screenshot: {source}")
        original = SCREEN_DIR / name
        doc_image = SCREEN_DIR / f"{Path(name).stem}_doc.png"
        copyfile(source, original)
        crop_for_doc(source, doc_image)
        copied[key] = {"original": original, "doc": doc_image}
    return copied


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 26, "00338D"),
        ("Heading 1", 18, "00338D"),
        ("Heading 2", 14, "1A1A2E"),
        ("Heading 3", 12, "1A1A2E"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_paragraph(doc, text, bold=False, color=None, size=10.5, align=None):
    paragraph = doc.add_paragraph()
    if align:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_font(run)


def add_image(doc, path, caption):
    if not Path(path).exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(9.7))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    set_font(cap_run, size=9, color="4A4A5A")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        shade_cell(hdr[idx], "E8EDF5")
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in hdr[idx].paragraphs:
            for run in paragraph.runs:
                set_font(run, bold=True, color="00338D")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=9.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LuminFlow Audit Management System")
    set_font(run, size=28, bold=True, color="00338D")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Consultant Review Pack / Platform Demo")
    set_font(run, size=18, bold=True, color="1E49E2")
    add_paragraph(doc, "Functional review, interactive experience, and constructive feedback collection for consulting team members", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_table(doc, ["Item", "Description"], [
        ["Platform Positioning", "A transparent collaboration platform for the audit domain, enhancing audit visualization, collaboration efficiency, and result credibility."],
        ["Demo Scope", "Dashboard, Audit Objectives, Sampling Preview, Impact Simulation, Team Management."],
        ["Screenshot Language", "Main pages use English UI; language switch, role switch, and AI assistant use partial screenshots for illustration."],
        ["Supporting Deliverable", "Standalone HTML interactive demo page: LuminFlow_Interactive_Demo.html."],
    ])
    doc.add_page_break()


def add_global_sections(doc, shots):
    doc.add_heading("1. Quick Tour", level=1)
    add_paragraph(doc, "LuminFlow integrates audit objectives, control testing, smart sampling, impact simulation, AI-assisted explanations, and team collaboration into a single transparent workbench. Consultants can use this document to quickly understand each page's features, core operational workflows, and audit workflow value.")
    add_table(doc, ["Page", "Core Purpose", "Consultant Review Focus"], [
        ["Dashboard", "Project cockpit displaying progress, risks, PBCs, pending items, and activity.", "Whether KPI definitions cover management and audit team needs."],
        ["Objective", "Audit objectives, control points, deficiencies, and COSO relationship graph.", "Whether node granularity, deficiency priorities, and evidence linkage are reasonable."],
        ["Sampling", "Risk-oriented sampling recommendations with AI logic explanations.", "Whether sampling rationale is auditable, defensible, and documentable."],
        ["Impact", "Cascading impact simulation and mitigation recommendations for change events.", "Whether impact chains and recommendations align with professional judgment."],
        ["Team", "Multi-role collaboration, AI approval, client interaction, and evidence checking.", "Whether permission boundaries and approval mechanisms meet quality control requirements."],
    ])

    doc.add_heading("2. Role & Permission Model", level=1)
    add_table(doc, ["Role", "Focus Areas", "Typical Visible Capabilities"], [
        ["CFO / Client Management", "Project status, risk summaries, PBC responses, client interactions.", "Simplified dashboard, client interaction statistics, approved AI summaries."],
        ["Auditor / Audit Manager", "Execution progress, control testing, sampling plans, evidence completeness.", "Full workbench, sampling parameters, AI approval queue, evidence checker."],
        ["Partner / Engagement Partner", "Quality review, significant matters, review comments, and risk judgments.", "Review KPIs, approval queue, QC perspective, and high-risk items."],
    ])
    add_image(doc, shots["role"]["doc"], "Role switcher partial screenshot: CFO, Audit Manager, Partner — three perspective entry points.")

    doc.add_heading("3. Global AI & Internationalization Capabilities", level=1)
    add_bullets(doc, [
        "Multi-language Support: Based on zustand state management and a custom useTranslation hook, default language is English, with toggle state persisted to localStorage.",
        "Full-stack i18n Coverage: Navigation, buttons, labels, table column names, mock data, AI-generated content, dynamic fields, and chart titles all support both English and Chinese.",
        "AI Assistant Mechanism: The frontend sends the current page route, user role, and conversation content as context to the backend, supporting page-relevant quick prompts and streaming responses.",
        "No-API-Key Fallback: When backend environment variables are not configured, a mock mode demonstrates AI capabilities and recommendation logic.",
    ])
    add_image(doc, shots["language"]["doc"], "Language switch partial screenshot: English main interface with one-click Chinese toggle.")
    add_image(doc, shots["ai"]["doc"], "AI Assistant partial screenshot: drawer-style conversation, role context tags, quick prompts, and input field.")
    doc.add_page_break()


def add_page_section(doc, page, shots):
    doc.add_heading(page["title"], level=1)
    add_image(doc, shots[page["screenshot"]]["doc"], f"{page['title']} main page screenshot")
    doc.add_heading("Page Objective", level=2)
    add_paragraph(doc, page["goal"])
    doc.add_heading("Core Modules", level=2)
    add_bullets(doc, page["modules"])
    doc.add_heading("Core Workflow", level=2)
    for idx, item in enumerate(page["flow"], 1):
        add_paragraph(doc, f"{idx}. {item}")
    doc.add_heading("Audit Workflow Value", level=2)
    add_paragraph(doc, page["value"])
    doc.add_heading("Consultant Feedback Suggestions", level=2)
    add_paragraph(doc, page["feedback"])
    doc.add_paragraph()


def add_impact_extra(doc, shots):
    doc.add_heading("Impact Simulation Supplementary Screenshots", level=2)
    add_image(doc, shots["impact_config"]["doc"], "Impact simulation configuration state: selecting event type, severity, impact dimensions, and response strategy.")
    add_image(doc, shots["impact_loading"]["doc"], "Impact simulation loading state: processing animation and network visualization during simulation execution.")


def add_summary(doc):
    doc.add_heading("Feature Highlights Summary", level=1)
    add_table(doc, ["Highlight", "Innovation Value", "Consultant Evaluation Question"], [
        ["Impact Simulation", "Transforms control deficiencies and business changes into visual impact chains.", "Does the propagation logic align with actual audit judgment?"],
        ["Smart Sampling", "Generates sample recommendations combining risk scores, confidence, deficiency history, and branch coverage.", "Are sampling explanations sufficient for workpaper inclusion?"],
        ["Audit Objective Graph", "Places audit objectives, controls, tests, and findings into a single relationship network.", "Does the node hierarchy align with project methodology?"],
        ["AI Contextual Assistant", "Generates explanations, recommendations, and communication summaries based on current page and role.", "Does AI output need more citations and evidence constraints?"],
        ["Multi-Role Transparent Collaboration", "Establishes approval and visibility boundaries between client transparency and quality control.", "Do permission boundaries meet independence and disclosure requirements?"],
    ])

    doc.add_heading("Improvement Suggestions Collection Form", level=1)
    add_table(doc, ["Feedback Dimension", "Consultant Comments", "Priority", "Follow-up Actions"], [
        ["Business Value", "", "", ""],
        ["Usability & Interaction", "", "", ""],
        ["Data Credibility", "", "", ""],
        ["AI Explainability", "", "", ""],
        ["Role Permissions", "", "", ""],
        ["Consultant Deliverable Fit", "", "", ""],
    ])


def build_doc(shots):
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_global_sections(doc, shots)
    for page in PAGES:
        add_page_section(doc, page, shots)
        if page["title"].startswith("Impact"):
            add_impact_extra(doc, shots)
        doc.add_page_break()
    add_summary(doc)
    doc.save(OUTPUT)


def main():
    ensure_dirs()
    shots = prepare_screenshots()
    build_doc(shots)
    print(f"DOCX generated: {OUTPUT}")
    print(f"Screenshots copied: {SCREEN_DIR}")


if __name__ == "__main__":
    main()

from pathlib import Path
from shutil import copyfile
from PIL import Image

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"e:\Hackathon\LuminFlow")
DEMO_DIR = ROOT / "docs" / "demo"
SCREEN_DIR = DEMO_DIR / "screenshots"
SOURCE_DIR = Path(r"C:\Program Files\Qoder")
OUTPUT = DEMO_DIR / "LuminFlow_Platform_Demo_For_Consultants_zh.docx"

SCREENSHOTS = {
    "dashboard": "luminflow-dashboard-zh.png",
    "objective": "luminflow-objective-zh.png",
    "sampling": "luminflow-sampling-zh.png",
    "impact_config": "luminflow-impact-config-zh.png",
    "impact_loading": "luminflow-impact-loading-zh.png",
    "impact_result": "luminflow-impact-result-zh.png",
    "team": "luminflow-team-zh.png",
    "role": "luminflow-role-switcher-zh.png",
    "ai": "luminflow-ai-drawer-zh.png",
    "language": "luminflow-language-zh.png",
}

PAGES = [
    {
        "title": "门户首页",
        "screenshot": "dashboard",
        "goal": "作为审计项目驾驶舱，集中呈现项目进度、关键指标、风险热区、控制状态、待办事项和活动时间线。",
        "modules": [
            "KPI 卡片：按角色过滤展示审计整体进度、控制覆盖率、PBC 完成率、缺陷数量和风险评分。",
            "风险热力图：基于 RCM 矩阵和 RAWTC 评分呈现风险控制热区，帮助快速定位高风险控制点。",
            "状态环形图 / 待处理事项：展示控制测试状态分布、PBC 逾期项和待处理事项。",
            "活动时间线：记录关键交付、审批、客户响应和复核活动。",
        ],
        "flow": [
            "进入门户首页后先查看项目总体健康度和关键风险指标。",
            "按角色切换 CFO、审计经理或合伙人视角，观察同一项目在不同权限下的摘要差异。",
            "从高风险 KPI、逾期 PBC 或风险热区进入审计目标、抽样预览或团队管理模块做进一步分析。",
        ],
        "value": "门户首页让管理层和项目团队以同一事实基础沟通进度与风险，减少审计状态汇报中的信息不对称。",
        "feedback": "请顾问重点评估 KPI 口径是否覆盖项目管理、质量复核和客户沟通三类核心需求。",
    },
    {
        "title": "审计目标",
        "screenshot": "objective",
        "goal": "以树状图和关联卡片展示审计目标、风险区域、控制点、测试程序和发现项之间的关系。",
        "modules": [
            "审计目标树：展示审计目标层级、控制点状态、缺陷节点和风险区域。",
            "节点详情：点击节点后查看控制说明、执行状态、风险评分、发现项和整改建议。",
            "雷达图 / 双进度条：展示 COSO 覆盖度、审计进度和客户侧响应进度。",
            "发现项表格 / AI 洞察面板：汇总缺陷严重度、偏差率、影响评估和 AI 优先级建议。",
        ],
        "flow": [
            "在搜索框中定位控制点或发现项，例如 CTRL-002 收入证明审核。",
            "点击节点打开详情，核查测试程序、COSO 映射和缺陷证据。",
            "结合雷达图、双进度条和发现项表格，确定下一步复核或整改优先级。",
        ],
        "value": "审计目标页把传统审计矩阵转化为可探索图谱，便于顾问检查测试范围完整性和控制缺陷传导关系。",
        "feedback": "请顾问关注节点粒度、COSO 映射和缺陷优先级是否符合实际审计方法论。",
    },
    {
        "title": "抽样预览",
        "screenshot": "sampling",
        "goal": "基于总体规模、风险评分、置信水平、上年缺陷和系统变更，生成可解释的智能抽样方案。",
        "modules": [
            "样本量卡片：展示推荐样本范围、总体规模、置信水平和期间。",
            "参数调节器：允许调整风险阈值、置信水平和高风险控制加权系数。",
            "旭日图 / 时间窗口卡片：展示文档类别分布和时间窗口分布。",
            "AI 逻辑解释器 / 样本明细表：解释抽样依据并展示样本明细。",
        ],
        "flow": [
            "顾问查看系统推荐样本范围和文档分布。",
            "调整置信水平或风险阈值后重新计算样本量。",
            "查看 AI 解释、历史趋势和样本明细，判断抽样方案是否可被复核和质询。",
        ],
        "value": "智能抽样使抽样理由从黑箱经验转为可解释、可调整、可复核的风险导向方案。",
        "feedback": "请顾问评估抽样参数、样本分布和解释文本是否足以支持底稿留痕。",
    },
    {
        "title": "影响模拟",
        "screenshot": "impact_result",
        "goal": "模拟控制缺陷、组织调整、法规变更或 IT 系统变更对审计范围、样本量、时间计划和风险指标的连锁影响。",
        "modules": [
            "事件配置：配置事件类型、严重度、影响维度和应对策略。",
            "力导向网络图：以网络图展示直接、间接和潜在影响节点。",
            "对比表：比较模拟前后的残余风险、控制有效性、合规指标、运营效率和成本影响。",
            "AI 建议：根据模拟结果生成缓释建议、优先级和风险降低幅度。",
        ],
        "flow": [
            "选择变更事件并设置严重度、影响维度和应对策略。",
            "点击运行模拟，等待网络节点传播动画和计算过程完成。",
            "查看影响链、指标变化和 AI 推荐，并决定是否调整审计范围或增加补偿性控制。",
        ],
        "value": "影响模拟将风险事件从静态描述转为动态决策依据，支持顾问提前评估变更对审计计划的影响。",
        "feedback": "请顾问重点验证影响维度、传播逻辑和 AI 建议是否符合审计专业判断。",
    },
    {
        "title": "团队管理",
        "screenshot": "team",
        "goal": "通过可见性管理、AI 审批、客户互动统计和证据检查，支持审计团队、客户和合伙人的透明协作。",
        "modules": [
            "可见性面板：控制客户可见、团队内部可见和合伙人复核可见的内容边界。",
            "审批队列：对 AI 生成内容进行审计经理审批和合伙人复核，避免未经确认的信息外发。",
            "互动图表 / 热点问题栏：追踪客户问题、热点关注和响应效率。",
            "证据检查器：检查关键证据缺失、逾期和复核状态。",
        ],
        "flow": [
            "审计经理查看团队协作和可见性状态。",
            "对 AI 生成摘要或客户回复执行审批，必要时提交合伙人复核。",
            "结合客户互动热点和证据检查结果，推动 PBC 补充和缺陷闭环。",
        ],
        "value": "团队管理页在提高透明度的同时保留审计质量控制和权限边界，降低客户沟通与 AI 输出风险。",
        "feedback": "请顾问评估审批颗粒度、证据检查规则和客户可见边界是否满足项目质量控制要求。",
    },
]


def ensure_dirs():
    DEMO_DIR.mkdir(parents=True, exist_ok=True)
    SCREEN_DIR.mkdir(parents=True, exist_ok=True)


def crop_for_doc(src: Path, dst: Path):
    with Image.open(src) as image:
        image = image.convert("RGB")
        width, height = image.size
        crop_height = min(height, 1100)
        cropped = image.crop((0, 0, width, crop_height))
        if cropped.width > 1800:
            ratio = 1800 / cropped.width
            cropped = cropped.resize((1800, int(cropped.height * ratio)), Image.Resampling.LANCZOS)
        cropped.save(dst, "PNG", optimize=True)


def prepare_screenshots():
    copied = {}
    for key, name in SCREENSHOTS.items():
        source = SOURCE_DIR / name
        if not source.exists():
            raise FileNotFoundError(f"Missing screenshot: {source}")
        original = SCREEN_DIR / name
        doc_image = SCREEN_DIR / f"{Path(name).stem}_doc.png"
        copyfile(source, original)
        crop_for_doc(source, doc_image)
        copied[key] = {"original": original, "doc": doc_image}
    return copied


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 26, "00338D"),
        ("Heading 1", 18, "00338D"),
        ("Heading 2", 14, "1A1A2E"),
        ("Heading 3", 12, "1A1A2E"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_paragraph(doc, text, bold=False, color=None, size=10.5, align=None):
    paragraph = doc.add_paragraph()
    if align:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_font(run)


def add_image(doc, path, caption):
    if not Path(path).exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(9.7))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    set_font(cap_run, size=9, color="4A4A5A")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        shade_cell(hdr[idx], "E8EDF5")
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in hdr[idx].paragraphs:
            for run in paragraph.runs:
                set_font(run, bold=True, color="00338D")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=9.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("明鉴 LuminFlow 审计管理系统")
    set_font(run, size=28, bold=True, color="00338D")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("顾问评审材料 / 平台演示文档")
    set_font(run, size=18, bold=True, color="1E49E2")
    add_paragraph(doc, "面向顾问团队的功能评审、互动体验和建设性反馈收集材料", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_table(doc, ["项目", "说明"], [
        ["平台定位", "面向审计领域的透明协作平台，提升审计可视化、协作效率与结果可信度。"],
        ["演示范围", "门户首页、审计目标、抽样预览、影响模拟、团队管理。"],
        ["截图语言", "主页面采用中文界面；语言切换、角色切换、AI 助手使用局部截图说明。"],
        ["配套交付", "独立 HTML 互动演示页：LuminFlow_Interactive_Demo_zh.html。"],
    ])
    doc.add_page_break()


def add_global_sections(doc, shots):
    doc.add_heading("1. 快速导览", level=1)
    add_paragraph(doc, "明鉴 LuminFlow 将审计目标、控制测试、智能抽样、影响模拟、AI 辅助解释和团队协作整合在同一个透明工作台中。顾问可通过本材料快速了解各页面功能、核心操作流程和审计工作流价值。")
    add_table(doc, ["页面", "核心用途", "顾问评审重点"], [
        ["门户首页", "项目驾驶舱，展示进度、风险、PBC、待办和活动。", "KPI 口径是否覆盖管理层和审计团队需求。"],
        ["审计目标", "审计目标、控制点、缺陷和 COSO 关系图谱。", "节点粒度、缺陷优先级和证据关联是否合理。"],
        ["抽样预览", "风险导向抽样建议和 AI 逻辑解释。", "抽样依据是否可复核、可质询、可留痕。"],
        ["影响模拟", "变更事件的连锁影响模拟和缓释建议。", "影响链和建议是否符合专业判断。"],
        ["团队管理", "多角色协作、AI 审批、客户互动和证据检查。", "权限边界和审批机制是否满足质量控制。"],
    ])

    doc.add_heading("2. 角色与权限模型", level=1)
    add_table(doc, ["角色", "关注重点", "典型可见能力"], [
        ["CFO / 客户管理层", "项目状态、风险摘要、PBC 响应、客户互动。", "简化仪表盘、客户互动统计、已批准 AI 摘要。"],
        ["审计员 / 审计经理", "执行进度、控制测试、抽样方案、证据完整性。", "完整工作台、抽样参数、AI 审批队列、证据检查。"],
        ["合伙人 / 项目合伙人", "质量复核、重大事项、复核意见和风险判断。", "复核类 KPI、审批复核、质量控制视角和高风险事项。"],
    ])
    add_image(doc, shots["role"]["doc"], "角色切换局部截图：CFO、审计经理、合伙人——三类视角入口。")

    doc.add_heading("3. 全局 AI 与国际化能力", level=1)
    add_bullets(doc, [
        "多语言支持：基于 zustand 状态管理和自研 useTranslation hook，默认语言为中文，切换状态持久化至 localStorage。",
        "全栈国际化覆盖：导航、按钮、标签、表格列名、Mock 数据、AI 生成内容、动态字段和图表标题均支持中英文。",
        "AI 助手机制：前端将当前页面路由、用户角色和对话内容作为上下文发送给后端，支持页面相关快捷问题和流式响应。",
        "无 API Key 降级：后端环境变量未配置时，可使用 Mock 模式展示 AI 能力和推荐逻辑。",
    ])
    add_image(doc, shots["language"]["doc"], "语言切换局部截图：中文主界面可一键切换英文。")
    add_image(doc, shots["ai"]["doc"], "AI 助手局部截图：抽屉式对话、角色上下文标签、快捷提示词和输入框。")
    doc.add_page_break()


def add_page_section(doc, page, shots):
    doc.add_heading(page["title"], level=1)
    add_image(doc, shots[page["screenshot"]]["doc"], f"{page['title']} 主页面截图")
    doc.add_heading("页面目标", level=2)
    add_paragraph(doc, page["goal"])
    doc.add_heading("核心功能模块", level=2)
    add_bullets(doc, page["modules"])
    doc.add_heading("核心操作流程", level=2)
    for idx, item in enumerate(page["flow"], 1):
        add_paragraph(doc, f"{idx}. {item}")
    doc.add_heading("审计工作流价值", level=2)
    add_paragraph(doc, page["value"])
    doc.add_heading("顾问反馈建议", level=2)
    add_paragraph(doc, page["feedback"])
    doc.add_paragraph()


def add_impact_extra(doc, shots):
    doc.add_heading("影响模拟过程补充截图", level=2)
    add_image(doc, shots["impact_config"]["doc"], "影响模拟配置态：选择事件类型、严重度、影响维度和应对策略。")
    add_image(doc, shots["impact_loading"]["doc"], "影响模拟加载态：运行模拟时展示处理动画与网络可视化。")


def add_summary(doc):
    doc.add_heading("功能亮点总结", level=1)
    add_table(doc, ["亮点", "创新价值", "顾问可评估问题"], [
        ["影响模拟", "将控制缺陷和业务变更转化为可视化影响链。", "传播逻辑是否符合实际审计判断？"],
        ["智能抽样", "结合风险评分、置信水平、缺陷历史和分支行覆盖生成样本建议。", "抽样解释是否足以进入工作底稿？"],
        ["审计目标图谱", "把审计目标、控制、测试和发现项放入同一关系网络。", "节点层级是否贴近项目方法论？"],
        ["AI 上下文助手", "基于当前页面和角色生成解释、建议和沟通摘要。", "AI 输出是否需要更多引用与证据约束？"],
        ["多角色透明协作", "在客户透明和质量控制之间建立审批与可见性边界。", "权限边界是否满足独立性与信息披露要求？"],
    ])

    doc.add_heading("改进建议收集表", level=1)
    add_table(doc, ["反馈维度", "顾问意见", "优先级", "后续动作"], [
        ["业务价值", "", "", ""],
        ["可用性与交互", "", "", ""],
        ["数据可信度", "", "", ""],
        ["AI 可解释性", "", "", ""],
        ["角色权限", "", "", ""],
        ["顾问交付适配", "", "", ""],
    ])


def build_doc(shots):
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_global_sections(doc, shots)
    for page in PAGES:
        add_page_section(doc, page, shots)
        if page["title"] == "影响模拟":
            add_impact_extra(doc, shots)
        doc.add_page_break()
    add_summary(doc)
    doc.save(OUTPUT)


def main():
    ensure_dirs()
    shots = prepare_screenshots()
    build_doc(shots)
    print(f"DOCX generated: {OUTPUT}")
    print(f"Screenshots copied: {SCREEN_DIR}")


if __name__ == "__main__":
    main()

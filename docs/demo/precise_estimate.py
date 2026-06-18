# -*- coding: utf-8 -*-
"""Corrected page count estimator using python-docx (all values in EMU -> pt)"""
from docx import Document
from docx.shared import Pt, Emu

doc = Document(r"e:\Hackathon\LuminFlow\docs\demo\LuminFlow_Hackathon_Submission.docx")
s = doc.sections[0]

# Page dimensions (python-docx returns EMU)
page_h_emu = s.page_height  # EMU
page_w_emu = s.page_width   # EMU
top_m = s.top_margin or 914400
bot_m = s.bottom_margin or 914400
left_m = s.left_margin or 914400
right_m = s.right_margin or 914400

usable_h_emu = page_h_emu - top_m - bot_m
usable_w_emu = page_w_emu - left_m - right_m
usable_h_pt = usable_h_emu / 12700
usable_w_pt = usable_w_emu / 12700

print(f"Page: {page_w_emu} x {page_h_emu} EMU = {page_w_emu/36000:.1f} x {page_h_emu/36000:.1f} mm")
print(f"Margins: {top_m/36000:.1f}mm each side (1 inch)")
print(f"Usable: {usable_w_pt:.0f} x {usable_h_pt:.0f} pt")
print()

total_pt = 0

for para in doc.paragraphs:
    txt = para.text.strip()
    style_name = para.style.name if para.style else 'Normal'
    
    # Determine actual font size from runs
    font_pt = None
    for run in para.runs:
        if run.font.size:
            font_pt = run.font.size.pt
            break
    
    if font_pt is None:
        # Fall back to style defaults
        if style_name == 'Title': font_pt = 24
        elif style_name == 'Heading 1': font_pt = 15
        elif style_name == 'Heading 2': font_pt = 12
        elif style_name == 'Heading 3': font_pt = 11
        else: font_pt = 10
    
    # Line spacing
    pf = para.paragraph_format
    line_spacing = 1.0
    try:
        if pf.line_spacing_rule == 3:  # EXACTLY
            line_spacing = pf.line_spacing / font_pt
        elif pf.line_spacing:
            if pf.line_spacing > 100:
                line_spacing = pf.line_spacing / 240  # twips->lines
            elif pf.line_spacing > 0:
                line_spacing = pf.line_spacing
    except:
        line_spacing = 1.0
    
    # Space before/after in pt (python-docx returns EMU)
    before_pt = (pf.space_before or 0) / 12700
    after_pt = (pf.space_after or 0) / 12700
    
    # Estimate lines
    if txt:
        char_w_pt = font_pt * 0.5  # approx avg char width
        chars_per_line = max(1, int(usable_w_pt / char_w_pt))
        num_lines = max(1, (len(txt) + chars_per_line - 1) // chars_per_line)
        para_pt = num_lines * font_pt * line_spacing + before_pt + after_pt
    else:
        para_pt = font_pt * line_spacing + before_pt + after_pt
    
    total_pt += para_pt

# Table (1 table with ~14 data rows + header)
for table in doc.tables:
    table_pt = 0
    for row in table.rows:
        row_max_lines = 1
        for cell in row.cells:
            cell_lines = max(1, (len(cell.text.strip()) + 35 - 1) // 35)
            row_max_lines = max(row_max_lines, cell_lines)
        table_pt += row_max_lines * 10 * 1.0 + 2  # 10pt, 1.0 spacing, 2pt padding
    total_pt += table_pt + 6  # header space
    print(f"Table: {len(table.rows)} rows, ~{table_pt:.0f} pt")

# Dashboard image
# Image extent: 5486400 x 3786097 EMU = 152.4mm x 105.2mm
img_h_pt = 3786097 / 12700  # ≈ 298 pt
total_pt += img_h_pt + 20  # image + caption
print(f"Dashboard image: ~{img_h_pt:.0f} pt")

pages = total_pt / usable_h_pt
print(f"\n{'='*50}")
print(f"Total content height: {total_pt:.0f} pt")
print(f"Usable page height: {usable_h_pt:.0f} pt")
print(f"Estimated pages: {pages:.1f}")
print(f"Conservative (rounded up): {int(pages + 0.99)} pages")

# Verification: key elements
print(f"\n--- Key Elements Checklist ---")
checks = [
    ("Assumptions (5 items)", "Transparency Demand is Structural"),
    ("Rationales (5 items)", "Three-Phase Architecture"),
    ("Dashboard screenshot", "Figure 1"),
    ("Pain points (4 items)", "Information is spread"),
    ("Business Value (3 items)", "Transforms audit from black-box"),
    ("Prototype Scope (4 items)", "Synthetic bank lending"),
    ("COSO framework", "COSO 2013"),
    ("Challenge 2 response", "Challenge 2"),
    ("Tech stack table", "React 18"),
    ("End marker", "End of Submission"),
]
for label, keyword in checks:
    found = any(keyword.lower() in p.text.lower() for p in doc.paragraphs)
    # Also check tables
    if not found:
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    if keyword.lower() in c.text.lower():
                        found = True
                        break
    print(f"  {'✓' if found else '✗'} {label}")
